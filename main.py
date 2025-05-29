from io import BytesIO
from fastapi import FastAPI, File, Request, UploadFile
import datetime
import os
from fastapi.responses import FileResponse, Response
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from docxtpl import DocxTemplate, RichText
from docxtpldvm import DVM

from docx import Document

# pip install htmldocx
# from htmldocx import HtmlToDocx
# pip install html-for-docx
from html4docxmixed import HtmlToDocx

# from deepseekcode import HtmlToDocx

# from htmldocxme import HtmlToDocx


# pip install Spire.Doc
# from spire.doc import Document, FileFormat
# from spire.doc.common import *

app = FastAPI()

app.add_middleware(
  CORSMiddleware,
  # allow_origins= ["http://172.16.200.180:9001"],
  allow_origins= ["*"],
  allow_credentials= True,
  allow_methods= ["*"],
  allow_headers= ["*"]
)

# @app.post("/items/")
# async def create_item(item: dict):
#     item_dict = item
#     print(item_dict)
#     if 'tax' in item_dict:
#         price_with_tax = item['price'] + item['tax']
#         item_dict.update({"price_with_tax": price_with_tax})
#     return item_dict

@app.get("/file/")
async def get_file(id: str):
    file_path = os.path.abspath(os.path.join('output', id+ '.docx'))
    # Check if the file exists
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    # Return the file as a response
    return FileResponse(file_path)

@app.post("/uploadfile/")
async def create_upload_file(file: UploadFile = File(...)):
    file_location = f"./templates/{file.filename}"
    with open(file_location, "wb") as f:
        f.write(await file.read())
    return {"info": f"file '{file.filename}' saved at '{file_location}'"}

@app.get("/get-ip/")
async def get_ip(request: Request):
    client_ip = request.client.host
    return {"client_ip": request.headers.get('X-Real-Ip')} # read nginx/default.conf

@app.get("/thbc/get-ip/")
async def get_ip(request: Request):
    client_ip = request.client.host
    return {"client_ip": client_ip} # for dev, test

from jinja2 import Environment, BaseLoader, Template

# Tạo hàm tùy chỉnh để chèn nội dung HTML vào tài liệu Word
def html_to_rich_text(html_content, _doc):
    desc_document = Document()
    new_parser = HtmlToDocx()
    new_parser.add_html_to_document(html_content , desc_document)
    buff = BytesIO()
    desc_document.save(buff)
    # buff.seek(0)
    return _doc.new_subdoc(buff)

    # document = Document()
    # # Add a section to the document
    # section = document.AddSection()
    # # Set the page margins to 72 points (72 points = 1 inch)
    # section.PageSetup.Margins.All = 72
    # # Add a paragraph to the section
    # paragraph = section.AddParagraph()
    # # Add the HTML string to the paragraph
    # paragraph.AppendHTML(html_content)
    # # Save the result document to a DOCX file
    # # document.SaveToFile("HtmlStringToDocx.docx", FileFormat.Docx2016)
    # # Or save the result document to a DOC file
    # document.SaveToFile("HtmlStringToDoc.doc", FileFormat.Doc)
    # # buff = BytesIO()
    # # document.SaveToFile(buff, FileFormat.Docx)
    # document.Close()

import subprocess
from io import BytesIO
import tempfile

def convert_html_to_docx(html_content, output_format="docx"):
    # Tạo một đối tượng BytesIO để lưu trữ dữ liệu đầu ra
    buff = BytesIO()

    # Tạo một file tạm thời để lưu nội dung HTML
    with tempfile.NamedTemporaryFile(suffix=".html", delete=False) as temp_file:
        temp_file.write(html_content.encode("utf-8"))
        temp_file_path = temp_file.name

    try:
        # Gọi unoconv để chuyển đổi HTML sang DOCX
        process = subprocess.Popen(
            ["/home/diepnn/API/utils-linux", temp_file_path],
            stdout=subprocess.PIPE,  # Lấy đầu ra từ stdout
            stderr=subprocess.PIPE   # Lấy thông báo lỗi (nếu có)
        )
        stdout, stderr = process.communicate()

        # Kiểm tra nếu có lỗi
        if process.returncode != 0:
            raise Exception(f"unoconv failed with error: {stderr.decode()}")

        # Ghi dữ liệu đầu ra vào BytesIO
        buff.write(stdout)
        buff.seek(0)  # Đặt con trỏ về đầu để đọc dữ liệu sau này

        return buff

    except Exception as e:
        print(f"Error during conversion: {e}")
        return None

    finally:
        # Xóa file tạm thời sau khi hoàn thành
        os.remove(temp_file_path)

# # Ví dụ sử dụng
# html_content = """
# <!DOCTYPE html>
# <html>
# <head>
#     <title>Sample HTML</title>
# </head>
# <body>
#     <h1>Hello, World!</h1>
#     <p>This is a sample HTML content.</p>
# </body>
# </html>
# """

# output_buffer = convert_html_to_docx(html_content)

# if output_buffer:
#     # Lưu dữ liệu từ BytesIO vào file (hoặc xử lý tiếp)
#     with open("output.docx", "wb") as f:
#         f.write(output_buffer.getvalue())
#     print("Conversion successful! File saved as output.docx")
# else:
#     print("Conversion failed.")

# iii= 0
# def html_to_rich_text(html_content, _doc):
#     global iii
#     gg= convert_html_to_docx(html_content)
#     desc_document = Document(gg)
#     iii= iii+ 1
#     with open("output_"+ str(iii)+ ".docx", "wb") as f:
#         f.write(gg.getbuffer())
#     buff = BytesIO()
#     desc_document.save(buff)
#     # buff.seek(0)
#     return _doc.new_subdoc(buff)

#     # return _doc.new_subdoc(convert_html_to_docx(html_content))

def find_column_with_token(data, token_value):
    for item in data:
        if item.get('token') == token_value:
            return item.get('column')
    return None
def find_atr_with_token(data, token_value):
    for item in data:
        if item.get('token') == token_value:
            return item.get('atr')
    return None
def find_token_with_column(data, column_value):
    for item in data:
        if item.get('column') == column_value:
            return item.get('token')
    return None
def find_duplicate_indices(arr):
    indices = []
    i = 0
    while i < len(arr):
        start_index = i
        while i < len(arr) - 1 and arr[i] == arr[i + 1]:
            i += 1
        if start_index != i:
            indices.append([start_index+1, i+1])
        i += 1
    return indices

import uuid
import traceback
@app.post("/docx/")
async def gen_docx(data: dict):
    try:
        item= data['content']
        context= {}
        template = "./templates/"+ data['file'].replace('/', '').replace('\\', '')
        doc = DocxTemplate(template)
        for key in item:
            el= item[key]
            if 'html' in el and el['html'] is True:
                context.update({el['token']: html_to_rich_text(el['data'], doc)})
            else:
                if 'merges' in el:
                    merges= el['merges']
                    dat= el['data']
                    if isinstance(dat, list):
                        tokens = [item["token"] for item in merges if "token" in item]
                        if len(tokens)> 0:
                            columns = [item["column"] for item in merges if "column" in item]
                            # print(tokens)
                            # print(columns)
                            prev_index= {}
                            if len(columns)== 0:
                                t= tokens[0]
                                prev_index.update({t: find_duplicate_indices(dat)})
                                # print(prev_index)
                                context.update({t: DVM(prev_index[t])})
                            else:
                                for t in tokens:
                                    c= find_column_with_token(merges, t)
                                    if c is not None:
                                        prev_index.update({t: find_duplicate_indices(
                                            [item[c] for item in dat if c in item]
                                        )})
                                        print(prev_index)
                                        context.update({t: DVM(prev_index[t])})
                elif 'merges_child' in el:
                    merges= el['merges_child']
                    dat= el['data']
                    if isinstance(dat, list):
                        tokens = [item["token"] for item in merges if "token" in item]
                        atr = [item["atr"] for item in merges if "atr" in item]
                        columns = [item["column"] for item in merges if "column" in item]   
                        for t in tokens:
                            c= find_column_with_token(merges, t)
                            for d_0 in dat:
                                d= d_0[find_atr_with_token(merges, t)]
                                # print('d:', d)
                                if c is not None:
                                    vm_1= find_duplicate_indices(
                                        [item[c] for item in d if c in item]
                                    )
                                    d_0.update({t: DVM(vm_1)})
                                    # print('d_0', d_0)
                context.update({el['token']: el['data']})
        # print(context)

        doc.render(context)
        # Tạo một buffer để lưu nội dung tài liệu
        buffer = BytesIO()
        doc.save(buffer)
        id= str(uuid.uuid4())
        doc.save('output/'+ id+ '.docx')
        buffer.seek(0)
        if data['url'] is True:
            return {
                'id': id
            }
        # Trả về nội dung tài liệu mà không cần lưu thành file
        return Response(
            buffer.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # # context.update({'diepnn': [(1,3),(4,6),(7,9)]})
        # return context
    except Exception as e:
        traceback_str = traceback.format_exc()
        return {"error": str(e), "traceback": traceback_str}

